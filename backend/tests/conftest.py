"""Shared pytest fixtures for the backend test suite.

Env vars required by Pydantic Settings are set at module load time so that
``app.config.settings`` instantiates successfully without a real .env file.
"""

import os

# Must appear before any app imports — Settings() is instantiated at import time.
os.environ.setdefault("OPENAI_API_KEY", "sk-test-key")
os.environ.setdefault("POSTGRES_URL", "postgresql+asyncpg://test:test@localhost:5432/test")

from pathlib import Path
from uuid import uuid4

import fitz  # PyMuPDF
import pytest
from docx import Document

from app.models.schemas import ParsedSection


# ---------------------------------------------------------------------------
# PDF fixtures
# ---------------------------------------------------------------------------


@pytest.fixture()
def pdf_two_pages(tmp_path: Path) -> Path:
    """A two-page text-only PDF."""
    doc = fitz.open()

    page1 = doc.new_page()
    page1.insert_text((50, 72), "This is the text on page one.")

    page2 = doc.new_page()
    page2.insert_text((50, 72), "This is the text on page two.")

    path = tmp_path / "two_pages.pdf"
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture()
def pdf_image_only(tmp_path: Path) -> Path:
    """A single-page PDF whose only content is a raster image (no text)."""
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)

    # Build a small red square as a PNG and embed it.
    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 80, 80), False)
    pix.set_rect(pix.irect, (220, 50, 50))
    img_bytes = pix.tobytes("png")
    page.insert_image(fitz.Rect(100, 100, 180, 180), stream=img_bytes)

    path = tmp_path / "image_only.pdf"
    doc.save(str(path))
    doc.close()
    return path


# ---------------------------------------------------------------------------
# DOCX fixtures
# ---------------------------------------------------------------------------


@pytest.fixture()
def docx_with_headings(tmp_path: Path) -> Path:
    """A DOCX with two Heading 1 sections, each containing one paragraph."""
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the introduction paragraph.")
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph("This is the conclusion paragraph.")

    path = tmp_path / "with_headings.docx"
    doc.save(str(path))
    return path


@pytest.fixture()
def docx_no_headings(tmp_path: Path) -> Path:
    """A DOCX with two body paragraphs and no headings."""
    doc = Document()
    doc.add_paragraph("First paragraph of the document.")
    doc.add_paragraph("Second paragraph of the document.")

    path = tmp_path / "no_headings.docx"
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------


@pytest.fixture()
def document_id() -> object:
    """A fresh UUID for use as a document_id in parser calls."""
    return uuid4()


# ---------------------------------------------------------------------------
# Chunker fixtures
# ---------------------------------------------------------------------------


@pytest.fixture()
def parsed_sections() -> list[ParsedSection]:
    """Three ParsedSections from the same document with representative metadata."""
    doc_id = uuid4()
    return [
        ParsedSection(
            document_id=doc_id,
            source_file="contract.pdf",
            page_number=1,
            section_header="Definitions",
            text="This section defines the key terms used in this agreement.",
            section_index=0,
        ),
        ParsedSection(
            document_id=doc_id,
            source_file="contract.pdf",
            page_number=2,
            section_header="Obligations",
            text="The parties agree to perform all obligations set forth herein.",
            section_index=1,
        ),
        ParsedSection(
            document_id=doc_id,
            source_file="contract.pdf",
            page_number=3,
            section_header=None,
            text="Signature block and execution date.",
            section_index=2,
        ),
    ]
